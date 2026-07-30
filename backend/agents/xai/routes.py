import json
import logging
import os
import uuid
from datetime import datetime
from typing import Dict, Any, List, Optional
import io
import csv

from fastapi import APIRouter, HTTPException, Body
from fastapi.responses import StreamingResponse
import requests

from database import get_db
from agents.history_helper import save_investigation

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/xai", tags=["Explainability Agent"])

GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

# Suppress imports errors by importing reportlab and docx within functions
# reportlab and docx are already in requirements.txt

@router.post("/explain")
async def generate_xai_explanation(payload: Dict[str, Any] = Body(...)):
    """
    Receives outputs from all ScamON AI agents or case ID, correlates findings,
    generates a unified forensic report, translates it into the selected language,
    and saves the generated XAI report to the database and Case Folder.
    """
    language = payload.get("language", "English")
    case_id = payload.get("case_id")
    
    db = get_db()
    resolved_case_id = case_id
    
    website = payload.get("website")
    email = payload.get("email")
    call = payload.get("call")
    live_call = payload.get("live_call")
    sms = payload.get("sms")
    threat_correlation = payload.get("threat_correlation")
    complaint = payload.get("complaint")
    visual_scam = payload.get("visual_scam")
    
    # If case_id is passed and we have DB, load evidence from the Case Folder!
    if case_id and db is not None:
        case_doc = db["cases"].find_one({"case_id": case_id})
        if case_doc:
            evidence = case_doc.get("evidence", {})
            if "website" in evidence and not website:
                website = evidence["website"].get("data")
            if "email" in evidence and not email:
                email = evidence["email"].get("data")
            if "call" in evidence and not call:
                call = evidence["call"].get("data")
            if "live_call" in evidence and not live_call:
                live_call = evidence["live_call"].get("data")
            if "sms" in evidence and not sms:
                sms = evidence["sms"].get("data")
            if "visual_scam" in evidence and not visual_scam:
                visual_scam = evidence["visual_scam"].get("data")
            if "threat_correlation" in evidence and not threat_correlation:
                threat_correlation = evidence["threat_correlation"].get("data")
            
            reports = case_doc.get("reports", {})
            if "complaint" in reports and not complaint:
                complaint = reports["complaint"]

    telemetry = {}
    if website: telemetry["website"] = website
    if email: telemetry["email"] = email
    if call: telemetry["call"] = call
    if live_call: telemetry["live_call"] = live_call
    if sms: telemetry["sms"] = sms
    if visual_scam: telemetry["visual_scam"] = visual_scam
    if threat_correlation: telemetry["threat_correlation"] = threat_correlation
    if complaint: telemetry["complaint"] = complaint

    if not telemetry:
        raise HTTPException(status_code=400, detail="No agent outputs provided for explainability analysis.")

    # Formulate Prompt
    prompt = f"""You are the Lead Cyber Security Forensic Explainability (XAI) Agent for the ScamON AI multi-agent platform.
Your task is to analyze the combined evidence from all active detection agents, correlate their findings, detect common threat indicators, explain why the investigation reached its conclusion, and present a final, unified report.

You MUST write the entire report, including the overall summary, findings bullet points, and recommendations, in the following target language: {language}.

Active Agent Output Telemetry:
{json.dumps(telemetry, indent=2)}

Please generate a structured report. You must respond strictly with a JSON object. Do not include markdown formatting or extra text. Use exactly this JSON structure:
{{
  "overall_summary": "A detailed, human-readable paragraph explaining why the platform reached its conclusion, in {language}.",
  "agents_used": ["List of agent names used in the scan, e.g. Website Investigation Agent, Email Investigation Agent"],
  "findings": {{
    "website": ["List of key bullet point findings for Website scan, in {language}"],
    "email": ["List of key bullet point findings for Email scan, in {language}"],
    "call": ["List of key bullet point findings for Call scan, in {language}"],
    "live_call": ["List of key bullet point findings for Live Call scan, in {language}"],
    "sms": ["List of key bullet point findings for SMS scan, in {language}"],
    "visual_scam": ["List of key bullet point findings for Visual Scam screenshot, in {language}"],
    "threat_correlation": ["List of findings for Correlation scan, in {language}"],
    "complaint": ["List of findings for Complaint scan, in {language}"]
  }},
  "overall_risk": {{
    "risk_score": 92, // Combined overall risk score (0-100)
    "threat_level": "HIGH", // SAFE | LOW | MEDIUM | HIGH | CRITICAL
    "confidence": 95 // Confidence percentage (0-100)
  }},
  "final_recommendations": [
    "List of direct, actionable recommendations in {language}."
  ]
}}
"""

    xai_report = {}
    if GROQ_API_KEY and not GROQ_API_KEY.startswith("your_"):
        try:
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            }
            body = {
                "model": "llama-3.1-8b-instant",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1,
                "response_format": {"type": "json_object"},
                "max_tokens": 1200,
            }
            response = requests.post(GROQ_API_URL, headers=headers, json=body, timeout=12)
            if response.status_code == 200:
                result = response.json()
                content = result["choices"][0]["message"]["content"]
                xai_report = json.loads(content)
            else:
                logger.warning(f"Groq API returned HTTP error: {response.status_code}")
        except Exception as err:
            logger.error(f"XAI LLM correlation failed: {err}")

    # Fallback if Groq fails or is not set
    if not xai_report:
        logger.info("Falling back to deterministic rule-based XAI report generation...")
        agents_used = []
        findings = {}
        risk_scores = []
        confidences = [90]
        recommendations = []

        if website:
            agents_used.append("Website Investigation Agent")
            score = website.get("risk_score", 0)
            risk_scores.append(score)
            verdict = website.get("verdict") or website.get("threat_type") or "Clean Site"
            findings["website"] = [
                f"Domain analyzed: {website.get('domain', {}).get('name', 'N/A')}",
                f"Domain Age: {website.get('domain', {}).get('age_days', 0)} days",
                f"Website Verdict: {verdict}"
            ]
            recommendations.append("Be cautious when interacting with newly registered domains.")

        if email:
            agents_used.append("Email Investigation Agent")
            score = email.get("risk_score", 0)
            risk_scores.append(score)
            findings["email"] = [
                f"Sender: {email.get('sender', 'N/A')}",
                f"SPF: {email.get('headers_analysis', {}).get('spf', 'N/A')}",
                f"DMARC: {email.get('headers_analysis', {}).get('dmarc', 'N/A')}"
            ]
            recommendations.append("Check email sender address and verify email headers alignment.")

        if call:
            agents_used.append("Call Analysis Agent")
            score = call.get("risk_score", 0)
            risk_scores.append(score)
            findings["call"] = [
                f"Caller: {call.get('caller', 'N/A')}",
                f"Verdict: {call.get('ai_analysis', {}).get('threat_category', 'General')}"
            ]
            recommendations.append("Hang up on any unsolicited calls asking for payment or passwords.")

        if live_call:
            agents_used.append("Live Call Detector")
            score = live_call.get("risk_score", 0)
            risk_scores.append(score)
            findings["live_call"] = [
                f"Live threat category: {live_call.get('category', 'N/A')}",
                f"Scam confidence: {live_call.get('confidence', 'N/A')}"
            ]
            recommendations.append("Do not share OTP or verification codes with callers.")

        if sms:
            agents_used.append("SMS Investigation Agent")
            score = sms.get("risk_score") or sms.get("analysis", {}).get("risk_score", 0)
            risk_scores.append(score)
            verdict = sms.get("analysis", {}).get("classification") or sms.get("verdict") or "Clean SMS"
            findings["sms"] = [
                f"Sender: {sms.get('sms', {}).get('sender', sms.get('sender', 'N/A'))}",
                f"Message Content: {sms.get('sms', {}).get('message', sms.get('message', 'N/A'))}",
                f"Verdict: {verdict}"
            ]
            recommendations.append("Do not reply to unsolicited messages or click links in SMS.")

        if visual_scam:
            agents_used.append("Visual Scam Investigation Agent")
            score = visual_scam.get("risk_score", 0)
            risk_scores.append(score)
            findings["visual_scam"] = [
                f"Image Type: {visual_scam.get('image_type', 'N/A')}",
                f"Extracted Entities: {', '.join([f'{k}: {v}' for k, v in visual_scam.get('entities', {}).items() if v])[:150]}",
                f"Scam Category: {visual_scam.get('scam_category', 'N/A')}",
                f"Visual Flags: {', '.join(visual_scam.get('visual_indicators', []))[:150]}"
            ]
            recommendations.extend(visual_scam.get("recommendations", []))

        if threat_correlation:
            agents_used.append("Threat Correlation Agent")
            score = threat_correlation.get("risk_score", 0)
            risk_scores.append(score)
            findings["threat_correlation"] = ["Correlated threat score calculated across multiple metrics."]

        if complaint:
            agents_used.append("Complaint Agent")
            findings["complaint"] = [
                f"Complaint ID: {complaint.get('complaint_id', 'N/A')}",
                f"Recipient: {complaint.get('recipient', 'N/A')}"
            ]

        overall_score = max(risk_scores) if risk_scores else 0
        threat_level = "SAFE"
        if overall_score >= 75: threat_level = "CRITICAL"
        elif overall_score >= 50: threat_level = "HIGH"
        elif overall_score >= 25: threat_level = "MEDIUM"
        elif overall_score >= 10: threat_level = "LOW"

        xai_report = {
            "overall_summary": f"System completed multi-agent cybersecurity synthesis check in {language}.",
            "agents_used": agents_used,
            "findings": findings,
            "overall_risk": {
                "risk_score": overall_score,
                "threat_level": threat_level,
                "confidence": 90
            },
            "final_recommendations": recommendations or ["Audit complete. Maintain standard cyber hygiene practices."]
        }

    # Calculate dynamic risk contributions based on actual agent risk scores
    total_input_scores = 0
    scores_dict = {}
    if website and "risk_score" in website:
        scores_dict["website"] = website["risk_score"]
        total_input_scores += website["risk_score"]
    if email and "risk_score" in email:
        scores_dict["email"] = email["risk_score"]
        total_input_scores += email["risk_score"]
    if call and "risk_score" in call:
        scores_dict["call"] = call["risk_score"]
        total_input_scores += call["risk_score"]
    if live_call and "risk_score" in live_call:
        scores_dict["live_call"] = live_call["risk_score"]
        total_input_scores += live_call["risk_score"]
    if sms:
        sms_score = sms.get("risk_score") or sms.get("analysis", {}).get("risk_score", 0)
        scores_dict["sms"] = sms_score
        total_input_scores += sms_score
    if threat_correlation and "risk_score" in threat_correlation:
        scores_dict["threat_correlation"] = threat_correlation["risk_score"]
        total_input_scores += threat_correlation["risk_score"]

    contributors = {}
    if total_input_scores > 0:
        for name, score in scores_dict.items():
            contributors[name] = round((score / total_input_scores) * 100)
    else:
        # Default even distribution for safety if scores are 0
        active_agents = list(scores_dict.keys())
        if active_agents:
            pct = round(100 / len(active_agents))
            for agent in active_agents:
                contributors[agent] = pct

    xai_report["risk_contributors"] = contributors

    # Save XAI log
    investigation_id = f"xai_{uuid.uuid4().hex[:8]}"
    xai_report["investigation_id"] = investigation_id
    xai_report["timestamp"] = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    xai_report["language"] = language
    xai_report["status"] = "completed"
    xai_report["voice_generated"] = False

    try:
        save_investigation(
            agent_type="xai",
            investigation_id=investigation_id,
            risk_score=xai_report["overall_risk"]["risk_score"],
            threat_level=xai_report["overall_risk"]["threat_level"],
            input_data="AI Explainability Audit",
            summary=xai_report["overall_summary"],
            full_report=xai_report,
            recommendation=", ".join(xai_report["final_recommendations"][:3]),
            case_id=resolved_case_id
        )
        if resolved_case_id and db is not None:
            db["cases"].update_one(
                {"case_id": resolved_case_id},
                {"$set": {
                    "reports.xai_summary": xai_report,
                    "status": "Analysis Completed",
                    "updated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
                }}
            )
    except Exception as db_err:
        logger.warning(f"Failed to auto-save XAI report to history: {db_err}")

    return xai_report


@router.post("/chat")
async def chat_with_xai_context(payload: Dict[str, Any] = Body(...)):
    """
    Handles interactive chat queries below the XAI summary,
    constraining the assistant to answer only using the report context.
    """
    report = payload.get("report")
    query = payload.get("query", "")

    if not report or not query:
        raise HTTPException(status_code=400, detail="XAI Report context and user query are both required.")

    prompt = f"""You are the ScamON AI Forensic Assistant.
Answer the user's question about the security investigation strictly using only the findings, summary, and recommendation context provided below.
If the answer cannot be determined from this context, state clearly that the evidence does not contain that information.

Investigation Report Context:
{json.dumps(report, indent=2)}

User Question: {query}
"""

    if GROQ_API_KEY and not GROQ_API_KEY.startswith("your_"):
        try:
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            }
            body = {
                "model": "llama-3.1-8b-instant",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.2,
                "max_tokens": 600,
            }
            response = requests.post(GROQ_API_URL, headers=headers, json=body, timeout=8)
            if response.status_code == 200:
                result = response.json()
                answer = result["choices"][0]["message"]["content"]
                return {"answer": answer}
            else:
                return {"answer": f"Forensic engine returned error: {response.text}"}
        except Exception as err:
            return {"answer": f"Unable to reach reasoning agent: {str(err)}"}

    return {"answer": "Reasoning engine offline. Please check your GROQ_API_KEY configuration."}


@router.get("/{id}/export/pdf")
async def export_xai_pdf(id: str):
    """Generates a professional PDF representing the XAI report."""
    db = get_db()
    report = {}
    if db is not None:
        try:
            doc_node = db["investigations"].find_one({"investigation_id": id})
            if doc_node:
                report = doc_node.get("full_report") or {}
        except Exception:
            pass
            
    if not report:
        report = {
            "overall_summary": "Offline/Fallback multi-agent synthesis report. Connection to MongoDB registry is currently offline or unavailable.",
            "status": "COMPLETED",
            "overall_risk": {
                "risk_score": 75,
                "threat_level": "HIGH",
                "confidence": 90
            },
            "risk_contributors": {
                "website": 40,
                "email": 60
            },
            "findings": {
                "website": ["Security warnings found on scanned registries."],
                "email": ["SPF configuration checks failed."]
            },
            "final_recommendations": [
                "Verify domain certificates.",
                "Review email DKIM/SPF credentials."
            ]
        }
    
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = []

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'PDFTitle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#00E676'),
        spaceAfter=15
    )
    section_style = ParagraphStyle(
        'PDFSection',
        parent=styles['Heading2'],
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#00A3FF'),
        spaceBefore=12,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'PDFBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#333333'),
        spaceAfter=8
    )

    story.append(Paragraph("ScamON AI Forensic Audit Report", title_style))
    story.append(Paragraph("Multi-Agent Explainability (XAI) Synthesis Log", styles['Heading3']))
    story.append(Spacer(1, 15))

    # Metadata table
    risk = report.get("overall_risk", {})
    metadata = [
        [Paragraph("<b>Investigation ID:</b>", body_style), Paragraph(id, body_style)],
        [Paragraph("<b>Timestamp:</b>", body_style), Paragraph(doc_node.get("timestamp", ""), body_style)],
        [Paragraph("<b>Report Language:</b>", body_style), Paragraph(report.get("language", "English"), body_style)],
        [Paragraph("<b>Threat Level:</b>", body_style), Paragraph(risk.get("threat_level", "SAFE"), body_style)],
        [Paragraph("<b>Risk Score:</b>", body_style), Paragraph(f"{risk.get('risk_score', 0)}/100", body_style)],
        [Paragraph("<b>Confidence:</b>", body_style), Paragraph(f"{risk.get('confidence', 90)}%", body_style)]
    ]

    t = Table(metadata, colWidths=[120, 400])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f5f5f5')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#dddddd')),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(t)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Investigation Summary", section_style))
    story.append(Paragraph(report.get("overall_summary", "No summary details provided."), body_style))
    story.append(Spacer(1, 10))

    # Agent Findings
    story.append(Paragraph("Correlated Agent Findings", section_style))
    findings = report.get("findings", {})
    for agent_name, agent_findings in findings.items():
        if agent_findings:
            story.append(Paragraph(f"<b>{agent_name.replace('_', ' ').title()}:</b>", body_style))
            for f in agent_findings:
                story.append(Paragraph(f"• {f}", body_style))
            story.append(Spacer(1, 5))

    story.append(Spacer(1, 10))
    story.append(Paragraph("Risk Score Contributions", section_style))
    contributors = report.get("risk_contributors", {})
    contrib_str = ", ".join([f"{k.replace('_', ' ').title()}: {v}%" for k, v in contributors.items()])
    story.append(Paragraph(contrib_str or "None calculated.", body_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Actionable Recommendations", section_style))
    recs = report.get("final_recommendations", [])
    for r in recs:
        story.append(Paragraph(f"✓ {r}", body_style))

    pdf.build(story)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=xai_{id}.pdf"}
    )


@router.get("/{id}/export/docx")
async def export_xai_docx(id: str):
    """Generates a Microsoft Word DOCX document representing the XAI report."""
    db = get_db()
    report = {}
    timestamp = "N/A"
    if db is not None:
        try:
            doc_node = db["investigations"].find_one({"investigation_id": id})
            if doc_node:
                report = doc_node.get("full_report") or {}
                timestamp = doc_node.get("timestamp", "N/A")
        except Exception:
            pass
            
    if not report:
        report = {
            "overall_summary": "Offline/Fallback multi-agent synthesis report. Connection to MongoDB registry is currently offline or unavailable.",
            "status": "COMPLETED",
            "overall_risk": {
                "risk_score": 75,
                "threat_level": "HIGH",
                "confidence": 90
            },
            "risk_contributors": {
                "website": 40,
                "email": 60
            },
            "findings": {
                "website": ["Security warnings found on scanned registries."],
                "email": ["SPF configuration checks failed."]
            },
            "final_recommendations": [
                "Verify domain certificates.",
                "Review email DKIM/SPF credentials."
            ]
        }
    
    import docx
    
    doc = docx.Document()
    
    doc.add_heading("ScamON AI Forensic Audit Report", 0)
    doc.add_heading("Multi-Agent Explainability (XAI) Synthesis Log", level=2)
    
    p = doc.add_paragraph()
    p.add_run(f"Investigation ID: {id}\n").bold = True
    p.add_run(f"Timestamp: {timestamp}\n")
    p.add_run(f"Report Language: {report.get('language', 'English')}\n")
    risk = report.get("overall_risk", {})
    p.add_run(f"Threat Level: {risk.get('threat_level', 'SAFE')} (Risk Score: {risk.get('risk_score', 0)}/100)\n")
    p.add_run(f"Confidence: {risk.get('confidence', 90)}%\n")

    doc.add_heading("Investigation Summary", level=1)
    doc.add_paragraph(report.get("overall_summary", "No summary details provided."))

    doc.add_heading("Correlated Agent Findings", level=1)
    findings = report.get("findings", {})
    for agent_name, agent_findings in findings.items():
        if agent_findings:
            doc.add_heading(agent_name.replace('_', ' ').title(), level=2)
            for f in agent_findings:
                doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("Risk Score Contributions", level=1)
    contributors = report.get("risk_contributors", {})
    contrib_p = doc.add_paragraph()
    for k, v in contributors.items():
        contrib_p.add_run(f"{k.replace('_', ' ').title()}: {v}%\n")

    doc.add_heading("Actionable Recommendations", level=1)
    recs = report.get("final_recommendations", [])
    for r in recs:
        doc.add_paragraph(r, style='List Bullet')

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=xai_{id}.docx"}
    )
