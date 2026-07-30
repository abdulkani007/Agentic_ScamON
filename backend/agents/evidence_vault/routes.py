import io
import logging
import zipfile
import json
import os
import requests
from datetime import datetime
from typing import Any, Dict, List, Optional
from fastapi import APIRouter, HTTPException, Query, status
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from bson import ObjectId

from database import get_db
from .agent import add_evidence_to_vault, generate_case_id, calculate_integrity_hash

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/evidence", tags=["Evidence Vault"])

def clean_mongodb_doc(doc: Any) -> Any:
    if isinstance(doc, dict):
        return {k: clean_mongodb_doc(v) for k, v in doc.items()}
    elif isinstance(doc, list):
        return [clean_mongodb_doc(x) for x in doc]
    elif isinstance(doc, ObjectId):
        return str(doc)
    else:
        return doc

# Pydantic Schemas
class EvidencePayload(BaseModel):
    case_id: Optional[str] = None
    agent_source: str
    evidence_data: Dict[str, Any]

class StatusUpdatePayload(BaseModel):
    status: str

# ----------------- FastAPI Endpoints -----------------

@router.post("/cases", status_code=201)
async def create_case_folder():
    """Initializes a brand new Case Folder with a sequential Case ID."""
    db = get_db()
    if db is None:
        # Fallback offline mode ID
        fallback_id = f"SCAMON-2026-{int(datetime.utcnow().timestamp()) % 1000000:06d}"
        return {"case_id": fallback_id, "status": "Open", "offline": True}
        
    try:
        case_id = generate_case_id(db)
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
        
        new_case = {
            "case_id": case_id,
            "user_id": "default_user",
            "created_at": timestamp,
            "updated_at": timestamp,
            "status": "Open",
            "overall_risk_score": 0,
            "overall_threat_level": "SAFE",
            "agents_used": [],
            "evidence": {},
            "reports": {}
        }
        db["cases"].insert_one(new_case)
        return {"case_id": case_id, "status": "Open"}
    except Exception as e:
        logger.error(f"Failed to create new case folder: {e}")
        raise HTTPException(status_code=500, detail="Database write error.")

@router.get("/cases/active")
async def get_active_case():
    """Retrieves the active (most recently updated and non-Closed) Case Folder."""
    db = get_db()
    if db is None:
        return {"case_id": "SCAMON-2026-000001", "status": "Open", "offline": True}
    try:
        active_case = db["cases"].find_one({"status": {"$ne": "Closed"}}, sort=[("updated_at", -1)])
        if active_case:
            return {"case_id": active_case["case_id"], "status": active_case.get("status", "Open")}
        fallback_id = "SCAMON-2026-000001"
        return {"case_id": fallback_id, "status": "Open"}
    except Exception as e:
        logger.error(f"Failed to fetch active case folder: {e}")
        return {"case_id": "SCAMON-2026-000001", "status": "Open"}

@router.get("/cases")
async def list_cases(
    status: Optional[str] = Query(None, description="Filter cases by status"),
    threat_level: Optional[str] = Query(None, description="Filter cases by overall threat level"),
    search: Optional[str] = Query(None, description="Filter cases by Case ID or content"),
    sort_by: str = Query("newest", description="Sort by: newest, oldest, highest_risk, lowest_risk")
):
    """Retrieves all Case Folders matching filter criteria."""
    db = get_db()
    if db is None:
        return []
        
    try:
        query = {}
        if status and status != "All":
            query["status"] = status
        if threat_level and threat_level != "All":
            query["overall_threat_level"] = threat_level.upper()
        if search:
            query["$or"] = [
                {"case_id": {"$regex": search, "$options": "i"}},
                {"agents_used": {"$regex": search, "$options": "i"}},
                {"status": {"$regex": search, "$options": "i"}}
            ]
            
        cursor = db["cases"].find(query)
        
        # Sort logic
        if sort_by == "newest":
            cursor = cursor.sort("created_at", -1)
        elif sort_by == "oldest":
            cursor = cursor.sort("created_at", 1)
        elif sort_by == "highest_risk":
            cursor = cursor.sort("overall_risk_score", -1)
        elif sort_by == "lowest_risk":
            cursor = cursor.sort("overall_risk_score", 1)
            
        cases = []
        for doc in cursor:
            cases.append(clean_mongodb_doc(doc))
        return cases
    except Exception as e:
        logger.error(f"Failed to list Case Folders: {e}")
        return []

@router.get("/cases/{id}")
async def get_case_details(id: str):
    """Retrieves a single Case Folder."""
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="Database offline.")
        
    doc = db["cases"].find_one({"case_id": id})
    if not doc:
        raise HTTPException(status_code=404, detail="Case Folder not found.")
        
    return clean_mongodb_doc(doc)

@router.delete("/cases/{id}")
async def delete_case_folder(id: str):
    """Deletes a Case Folder."""
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="Database offline.")
        
    res = db["cases"].delete_one({"case_id": id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Case not found.")
    return {"message": "Case deleted successfully."}

@router.post("/cases/{id}/status")
async def update_case_status(id: str, payload: StatusUpdatePayload):
    """Updates the status of a Case Folder."""
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="Database offline.")
        
    timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    res = db["cases"].update_one(
        {"case_id": id},
        {"$set": {"status": payload.status, "updated_at": timestamp}}
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Case Folder not found.")
    return {"message": f"Case status updated to {payload.status}."}

@router.post("/add")
async def add_evidence(payload: EvidencePayload):
    """Endpoint for finished agents to push evidence to the vault."""
    db = get_db()
    if db is None:
        # Fallback offline mock response
        return {"message": "Saved (Offline Fallback)", "case_id": payload.case_id or "SCAMON-2026-000001"}
        
    try:
        updated_case = add_evidence_to_vault(
            db=db,
            case_id=payload.case_id,
            agent_source=payload.agent_source,
            evidence_data=payload.evidence_data
        )
        return clean_mongodb_doc(updated_case)
    except Exception as e:
        logger.error(f"Failed to add evidence: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ----------------- Document Exporters Helpers -----------------

def generate_pdf_stream(case_id: str, case_data: Dict[str, Any]) -> io.BytesIO:
    """Helper to generate a ReportLab PDF document for a Case Folder."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = []

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'VaultTitle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#00E676'),
        spaceAfter=10
    )
    section_style = ParagraphStyle(
        'VaultSection',
        parent=styles['Heading2'],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#00E676'),
        spaceBefore=15,
        spaceAfter=6,
        borderPadding=2
    )
    body_style = ParagraphStyle(
        'VaultBody',
        parent=styles['Normal'],
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#e2e8f0')
    )
    code_style = ParagraphStyle(
        'VaultCode',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#a7f3d0')
    )

    story.append(Paragraph("ScamON AI Digital Forensics Evidence Case", title_style))
    story.append(Paragraph(f"Case ID: {case_id} // Registry Status: {case_data.get('status', 'Open')}", body_style))
    story.append(Spacer(1, 10))

    # Meta Info table
    meta_data = [
        [Paragraph("<b>Risk Score</b>", body_style), Paragraph(str(case_data.get('overall_risk_score', 0)), body_style)],
        [Paragraph("<b>Threat Level</b>", body_style), Paragraph(case_data.get('overall_threat_level', 'SAFE'), body_style)],
        [Paragraph("<b>Investigation Date</b>", body_style), Paragraph(case_data.get('created_at', ''), body_style)],
        [Paragraph("<b>Last Modified</b>", body_style), Paragraph(case_data.get('updated_at', ''), body_style)]
    ]
    t = Table(meta_data, colWidths=[150, 350])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#0a192f')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#00E676')),
        ('PADDING', (0,0), (-1,-1), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#1e293b'))
    ]))
    story.append(t)
    story.append(Spacer(1, 15))

    # Evidence details
    story.append(Paragraph("Verifiable Forensics Evidence Logs", section_style))
    evidence = case_data.get("evidence", {})
    if not evidence:
        story.append(Paragraph("No evidence items cataloged yet.", body_style))
    else:
        for source, ev in evidence.items():
            story.append(Paragraph(f"▶ {ev.get('generated_by', source.upper())}", section_style))
            story.append(Paragraph(f"Timestamp: {ev.get('creation_time', '')} | Integrity Hash: {ev.get('integrity_hash', '')}", body_style))
            story.append(Spacer(1, 4))
            
            raw_data = ev.get("data") or {}
            dumped = json.dumps(raw_data, indent=2)
            story.append(Paragraph(dumped.replace('\n', '<br/>').replace(' ', '&nbsp;'), code_style))
            story.append(Spacer(1, 10))

    pdf.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_stream(case_id: str, case_data: Dict[str, Any]) -> io.BytesIO:
    """Helper to generate a Word DOCX document for a Case Folder."""
    import docx
    doc = docx.Document()
    
    doc.add_heading("ScamON AI - Digital Forensics Vault Case File", 0)
    doc.add_heading(f"Case Folder: {case_id}", level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"Status: {case_data.get('status', 'Open')}\n").bold = True
    p.add_run(f"Risk Score: {case_data.get('overall_risk_score', 0)}/100 ({case_data.get('overall_threat_level', 'SAFE')})\n")
    p.add_run(f"Created At: {case_data.get('created_at', '')}\n")
    p.add_run(f"Last Modified: {case_data.get('updated_at', '')}\n")
    
    doc.add_heading("Evidence Vault Registry Logs", level=2)
    evidence = case_data.get("evidence", {})
    if not evidence:
        doc.add_paragraph("No evidence logs collected yet.")
    else:
        for source, ev in evidence.items():
            doc.add_heading(ev.get("generated_by", source.upper()), level=3)
            doc.add_paragraph(f"Timestamp: {ev.get('creation_time', '')}\nIntegrity Hash: {ev.get('integrity_hash', '')}")
            
            # Add raw data dump
            raw_data = ev.get("data") or {}
            doc.add_paragraph(json.dumps(raw_data, indent=2))
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ----------------- Exporter Routes -----------------

@router.get("/cases/{id}/export/pdf")
async def export_case_pdf(id: str):
    """Streams case report PDF."""
    db = get_db()
    case_data = {}
    if db is not None:
        case_data = db["cases"].find_one({"case_id": id}) or {}
        case_data = clean_mongodb_doc(case_data)
        
    if not case_data:
        case_data = {
            "case_id": id,
            "status": "Offline / Not Found",
            "overall_risk_score": 0,
            "overall_threat_level": "UNKNOWN",
            "evidence": {}
        }
        
    pdf_stream = generate_pdf_stream(id, case_data)
    return StreamingResponse(
        pdf_stream,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=case_{id}.pdf"}
    )

@router.get("/cases/{id}/export/docx")
async def export_case_docx(id: str):
    """Streams case report Word DOCX document."""
    db = get_db()
    case_data = {}
    if db is not None:
        case_data = db["cases"].find_one({"case_id": id}) or {}
        case_data = clean_mongodb_doc(case_data)
        
    if not case_data:
        case_data = {
            "case_id": id,
            "status": "Offline / Not Found",
            "overall_risk_score": 0,
            "overall_threat_level": "UNKNOWN",
            "evidence": {}
        }
        
    docx_stream = generate_docx_stream(id, case_data)
    return StreamingResponse(
        docx_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=case_{id}.docx"}
    )

@router.get("/cases/{id}/export/json")
async def export_case_json(id: str):
    """Downloads Case Folder raw JSON data."""
    db = get_db()
    case_data = {}
    if db is not None:
        case_data = db["cases"].find_one({"case_id": id}) or {}
        
    if not case_data:
        raise HTTPException(status_code=404, detail="Case Folder not found.")
        
    if "_id" in case_data:
        case_data["_id"] = str(case_data["_id"])
        
    return JSONResponse(content=case_data)

@router.get("/cases/{id}/export/zip")
async def export_case_zip(id: str):
    """Streams a compiled .zip folder containing PDF, DOCX, and raw JSON logs."""
    db = get_db()
    case_data = {}
    if db is not None:
        case_data = db["cases"].find_one({"case_id": id}) or {}
        case_data = clean_mongodb_doc(case_data)
        
    if not case_data:
        raise HTTPException(status_code=404, detail="Case Folder not found.")
        
    if "_id" in case_data:
        case_data["_id"] = str(case_data["_id"])

    # Create zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        # Write raw JSON details
        json_data = json.dumps(case_data, indent=2, default=str)
        zip_file.writestr("case_details.json", json_data)
        
        # Write PDF report
        pdf_stream = generate_pdf_stream(id, case_data)
        zip_file.writestr("case_report.pdf", pdf_stream.getvalue())
        
        # Write DOCX report
        docx_stream = generate_docx_stream(id, case_data)
        zip_file.writestr("case_report.docx", docx_stream.getvalue())
        
        # Write individual agent json payloads
        evidence = case_data.get("evidence") or {}
        for source, ev in evidence.items():
            ev_data = ev.get("data") or {}
            zip_file.writestr(f"evidence_{source}.json", json.dumps(ev_data, indent=2, default=str))

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/x-zip-compressed",
        headers={"Content-Disposition": f"attachment; filename=case_{id}.zip"}
    )


GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

class VaultChatPayload(BaseModel):
    message: str
    history: List[Dict[str, str]] = []

@router.post("/chat")
async def chat_with_vault(payload: VaultChatPayload):
    """
    Handles AI Copilot queries for case search and evidence retrieval in the Digital Evidence Vault.
    """
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="Database connection offline.")
        
    import re
    # Extract domains or case IDs
    words = payload.message.lower().split()
    domains = []
    for w in words:
        match = re.search(r'([a-z0-9\-]+\.[a-z0-9\-.]+)', w)
        if match:
            domains.append(match.group(1).strip(".,!?;:"))

    # Extract agent keywords
    agent_keywords = []
    for k in ["sms", "email", "call", "website", "correlation"]:
        if k in payload.message.lower():
            agent_keywords.append(k)

    # Build query
    query = {}
    or_clauses = []

    if domains:
        for dom in domains:
            or_clauses.append({"evidence.website.data.url": {"$regex": dom, "$options": "i"}})
            or_clauses.append({"evidence.website.data.domain.name": {"$regex": dom, "$options": "i"}})
            or_clauses.append({"evidence.email.data.links_analysis.url": {"$regex": dom, "$options": "i"}})
            or_clauses.append({"evidence.email.data.links_analysis.domain": {"$regex": dom, "$options": "i"}})
            or_clauses.append({"evidence.email.data.sender": {"$regex": dom, "$options": "i"}})
            or_clauses.append({"case_id": {"$regex": dom, "$options": "i"}})

    if agent_keywords:
        for agent in agent_keywords:
            if agent == "correlation":
                or_clauses.append({"evidence.threat_correlation": {"$exists": True}})
            elif agent == "call":
                or_clauses.append({"evidence.call": {"$exists": True}})
                or_clauses.append({"evidence.live_call": {"$exists": True}})
            else:
                or_clauses.append({f"evidence.{agent}": {"$exists": True}})

    if or_clauses:
        query["$or"] = or_clauses

    try:
        if query:
            cases_cursor = db["cases"].find(query, {"_id": 0}).sort("updated_at", -1).limit(10)
        else:
            # Fallback to the 5 most recent cases
            cases_cursor = db["cases"].find({}, {"_id": 0}).sort("updated_at", -1).limit(5)
        cases_list = list(cases_cursor)
    except Exception as e:
        logger.error(f"Failed to fetch cases for vault chat: {e}")
        cases_list = []

    # Prepare a minimal summary of cases to prevent context window bloat while preserving key evidence indicators
    cases_summary = []
    for c in cases_list:
        evidence_summary = {}
        if "evidence" in c and c["evidence"]:
            for source, ev in c["evidence"].items():
                ev_data = ev.get("data") or {}
                # Extract key identifiers like domain, sender, phone numbers, risk rating
                if source == "website":
                    domain_val = ev_data.get("domain", "")
                    domain_name = domain_val if isinstance(domain_val, str) else domain_val.get("name", "")
                    evidence_summary["website"] = {
                        "url": ev_data.get("url", ""),
                        "domain": domain_name,
                        "verdict": ev_data.get("verdict", ""),
                        "risk_score": ev_data.get("risk_score", 0)
                    }
                elif source == "email":
                    evidence_summary["email"] = {
                        "subject": ev_data.get("subject", ""),
                        "sender": ev_data.get("sender", ""),
                        "receiver": ev_data.get("receiver", ""),
                        "risk_score": ev_data.get("risk_score", 0),
                        "llm_classification": ev_data.get("llm_classification", "")
                    }
                elif source == "call" or source == "live_call":
                    ai_analysis_val = ev_data.get("ai_analysis") or {}
                    rec_val = ai_analysis_val.get("recommended_action") if isinstance(ai_analysis_val, dict) else ""
                    evidence_summary[source] = {
                        "caller": ev_data.get("caller_phone") or ev_data.get("caller") or "",
                        "risk_score": ev_data.get("risk_score", 0),
                        "recommendation": rec_val or ev_data.get("recommendation", "")
                    }
                elif source == "sms":
                    analysis_val = ev_data.get("analysis") or {}
                    class_val = analysis_val.get("classification") if isinstance(analysis_val, dict) else ""
                    evidence_summary["sms"] = {
                        "sender": ev_data.get("sender", ""),
                        "classification": class_val or ev_data.get("verdict", ""),
                        "risk_score": ev_data.get("risk_score", 0)
                    }
                else:
                    evidence_summary[source] = {
                        "risk_score": ev_data.get("risk_score") or ev_data.get("overall_risk_score") or 0
                    }

        cases_summary.append({
            "case_id": c.get("case_id"),
            "status": c.get("status"),
            "overall_risk_score": c.get("overall_risk_score", 0),
            "overall_threat_level": c.get("overall_threat_level", "SAFE"),
            "created_at": c.get("created_at"),
            "agents_used": c.get("agents_used", []),
            "evidence": evidence_summary
        })

    system_prompt = """You are the ScamON SOC Evidence Vault Copilot, an expert AI cybersecurity assistant specialized in digital forensics and security operations.
Your job is to help analysts inspect, search, and locate evidence records from the Case Folders Directory stored in the Digital Evidence Vault.

You have access to the current cases index:
""" + json.dumps(cases_summary, indent=2) + """

Instructions:
1. When asked about specific evidence (e.g. "I need the evidence of youtube.com", or "evidence for sms"), search the cases list.
2. If matching cases are found:
   - Provide a clear, structured list of the matches.
   - For each matching case, specify the Case ID (e.g. `SCAMON-2026-000002`), the overall threat level, the risk score, the date created, and what specific evidence was found.
   - Guide the user on how they can view it in the dashboard (e.g., clicking on that Case ID in the left directory list, and clicking "View Evidence" under the corresponding agent card).
3. If no matching case or evidence is found, explain politely that no scan telemetry matches their search criteria.
4. Keep your responses concise, professional, and formatted in clean markdown.
5. If the user asks general questions about ScamON or the Evidence Vault, answer professionally.
"""

    messages = [{"role": "system", "content": system_prompt}]
    for h in payload.history:
        messages.append({"role": h.get("role", "user"), "content": h.get("content", "")})
    messages.append({"role": "user", "content": payload.message})

    if GROQ_API_KEY and not GROQ_API_KEY.startswith("your_"):
        try:
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            }
            body = {
                "model": "llama-3.1-8b-instant",
                "messages": messages,
                "temperature": 0.2,
                "max_tokens": 800,
            }
            response = requests.post(GROQ_API_URL, headers=headers, json=body, timeout=10)
            if response.status_code == 200:
                result = response.json()
                answer = result["choices"][0]["message"]["content"]
                return {"answer": answer}
            else:
                return {"answer": f"Forensic engine returned error: {response.text}"}
        except Exception as err:
            return {"answer": f"Unable to reach reasoning agent: {str(err)}"}
            
    return {"answer": "Reasoning engine offline. Please check your GROQ_API_KEY configuration."}

