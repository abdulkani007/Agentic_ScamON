import os
import uuid
import logging
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime
from typing import List, Dict, Any

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from database import get_db

logger = logging.getLogger(__name__)


def generate_complaint_package(report: Dict[str, Any], static_dir: str) -> Dict[str, Any]:
    """Generates professional legal complaint PDFs and Word documents from the investigation report.

    Saves files under static/complaints/{complaint_id}/
    """
    complaint_id = str(uuid.uuid4())
    complaint_dir = os.path.join(static_dir, "complaints", complaint_id)
    os.makedirs(complaint_dir, exist_ok=True)

    timestamp_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")
    is_call = "transcript" in report or "llm_analysis" in report

    # 1. Compile file names and relative paths
    files = {
        "complaint_pdf": f"/static/complaints/{complaint_id}/Complaint.pdf",
        "report_pdf": f"/static/complaints/{complaint_id}/Investigation_Report.pdf",
        "complaint_docx": f"/static/complaints/{complaint_id}/Complaint.docx",
    }

    complaint_pdf_path = os.path.join(complaint_dir, "Complaint.pdf")
    report_pdf_path = os.path.join(complaint_dir, "Investigation_Report.pdf")
    complaint_docx_path = os.path.join(complaint_dir, "Complaint.docx")

    # Call Transcript PDF (Call Scan only)
    call_transcript_pdf_path = None
    if is_call:
        files["transcript_pdf"] = f"/static/complaints/{complaint_id}/Call_Transcript.pdf"
        call_transcript_pdf_path = os.path.join(complaint_dir, "Call_Transcript.pdf")

    # Evidence Summary PDF (Website Scan only)
    evidence_pdf_path = None
    if not is_call:
        files["evidence_pdf"] = f"/static/complaints/{complaint_id}/Evidence_Report.pdf"
        evidence_pdf_path = os.path.join(complaint_dir, "Evidence_Report.pdf")

    # 2. Build Stylesheets
    styles = getSampleStyleSheet()
    
    # Custom Title & Header styles
    title_style = ParagraphStyle(
        'CyberTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=colors.HexColor('#007A3E'),
        spaceAfter=15,
        alignment=1 # Center
    )
    
    section_style = ParagraphStyle(
        'CyberSection',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        textColor=colors.HexColor('#0055A5'),
        spaceBefore=12,
        spaceAfter=8
    )

    body_style = ParagraphStyle(
        'CyberBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        textColor=colors.HexColor('#1a202c'),
        leading=14,
        spaceAfter=10
    )

    meta_label_style = ParagraphStyle(
        'CyberMetaLabel',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        textColor=colors.HexColor('#a0aec0'),
        leading=12
    )

    meta_val_style = ParagraphStyle(
        'CyberMetaVal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        textColor=colors.HexColor('#ffffff'),
        leading=12
    )

    # 3. Compile meaningful text variables
    risk_score = report.get("risk_score", 0)
    target_url = report.get("url", "N/A")
    threat_level = report.get("threat_type", report.get("ai_reasoning", {}).get("final_decision", "SUSPICIOUS"))
    
    sms_sender = report.get("sms_sender")
    sms_message = report.get("sms_message")
    sms_risk_score = report.get("sms_risk_score")
    sms_recommendation = report.get("sms_recommendation")
    visual_scam = report.get("visual_scam")
    
    # Build Complaint PDF
    try:
        doc = SimpleDocTemplate(complaint_pdf_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
        elements = []
        
        # Cyber Banner background simulation (translucent dark flowables)
        elements.append(Paragraph("<b>ScamON AI • CYBERCRIME COMPLAINT FILING</b>", title_style))
        elements.append(Spacer(1, 10))
        
        # Metadata block
        meta_data = [
            [Paragraph("Complaint ID:", meta_label_style), Paragraph(complaint_id, meta_val_style)],
            [Paragraph("Filing Date:", meta_label_style), Paragraph(timestamp_str, meta_val_style)],
            [Paragraph("Scam Target:", meta_label_style), Paragraph(target_url, meta_val_style)],
            [Paragraph("Risk Index:", meta_label_style), Paragraph(f"{risk_score}% RISK (Verdict: {threat_level})", meta_val_style)]
        ]
        t = Table(meta_data, colWidths=[100, 400])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#040f22')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING', (0,0), (-1,-1), 10),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#00E676')),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 20))
        
        # Complaint Letter body
        elements.append(Paragraph("<b>Subject: Complaint Regarding Suspected Cyber Scam</b>", section_style))
        elements.append(Spacer(1, 8))
        
        letter_body = f"""Dear Sir/Madam,
<br/><br/>
I would like to report a suspected cyber scam detected by the ScamON AI Multi-Agent Cybersecurity Platform.
<br/><br/>
Our autonomous threat auditing has flagged active malicious infrastructure targeting users at the following coordinates:<br/>
<b>Scam Target Identifier:</b> {target_url}<br/>
<b>Risk Probability Index:</b> {risk_score}%<br/>
<b>Identified Threat Pattern:</b> {threat_level}<br/>
<b>Detection Timestamp:</b> {timestamp_str}"""

        if sms_sender and sms_message:
            letter_body += f"""<br/><br/>
<b>SMS Phishing / SMiShing Telemetry:</b><br/>
• <b>SMS Sender:</b> {sms_sender}<br/>
• <b>SMS Message Content:</b> {sms_message}<br/>
• <b>SMS Risk Index:</b> {sms_risk_score}%<br/>
• <b>Forensics Recommendation:</b> {sms_recommendation}"""

        if visual_scam:
            letter_body += f"""<br/><br/>
<b>Visual Scam Investigation Telemetry:</b><br/>
• <b>Image Type:</b> {visual_scam.get('image_type', 'N/A')}<br/>
• <b>Risk Score:</b> {visual_scam.get('risk_score', 0)}% ({visual_scam.get('threat_level', 'LOW')})<br/>
• <b>OCR Text Snippet:</b> {visual_scam.get('extracted_text', '')[:150]}...<br/>
• <b>Scam Category:</b> {visual_scam.get('scam_category', 'N/A')}<br/>
• <b>Forensics Recommendation:</b> {', '.join(visual_scam.get('recommendations', []))[:150]}"""

        letter_body += f"""<br/><br/>
<b>Scam Indicators & Evidence Log:</b><br/>
The target has been flagged for multiple suspicious behaviors, including typosquatting brand similarities, invalid/self-signed SSL certificates, domain registration age anomalies, and/or social engineering patterns detected within communication transcripts. 
<br/><br/>
Kindly find the attached forensic investigation reports and supporting technical evidence for further evaluation and prompt action.
<br/><br/>
Regards,
<br/>
<b>ScamON AI Compliance Network</b>"""
        
        elements.append(Paragraph(letter_body, body_style))
        doc.build(elements)
    except Exception as e:
        logger.error(f"Failed to generate Complaint PDF: {e}")

    # Build Investigation Report PDF
    try:
        doc = SimpleDocTemplate(report_pdf_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
        elements = []
        
        elements.append(Paragraph("<b>FORENSIC INCIDENT INVESTIGATION REPORT</b>", title_style))
        elements.append(Spacer(1, 10))
        
        # Details Table
        report_data = [
            [Paragraph("Incident ID:", meta_label_style), Paragraph(report.get("investigation_id", "N/A"), meta_val_style)],
            [Paragraph("Filing Timestamp:", meta_label_style), Paragraph(timestamp_str, meta_val_style)],
            [Paragraph("Target URL:", meta_label_style), Paragraph(target_url, meta_val_style)],
            [Paragraph("Risk Score:", meta_label_style), Paragraph(f"{risk_score}/100", meta_val_style)],
            [Paragraph("Verdict Decision:", meta_label_style), Paragraph(threat_level, meta_val_style)],
            [Paragraph("Recommendation:", meta_label_style), Paragraph(report.get("recommendation", "N/A"), meta_val_style)]
        ]
        t = Table(report_data, colWidths=[120, 380])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#020b18')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#1d2f44')),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('LEFTPADDING', (0,0), (-1,-1), 8),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 15))
        
        elements.append(Paragraph("<b>Explainable AI (XAI) SOC Findings</b>", section_style))
        ai_summary = report.get("ai_reasoning", {}).get("summary", "Forensic threat investigation completed cleanly.")
        elements.append(Paragraph(ai_summary, body_style))
        
        elements.append(Paragraph("<b>Threat Auditing Logic Steps:</b>", section_style))
        steps = report.get("ai_reasoning", {}).get("reasoning_steps", [])
        if not steps:
            steps = ["All primary security modules verified without fatal exceptions."]
        for step in steps:
            elements.append(Paragraph(f"• {step}", body_style))
            
        doc.build(elements)
    except Exception as e:
        logger.error(f"Failed to generate Investigation Report PDF: {e}")

    # Build Call Transcript PDF if applicable
    if is_call and call_transcript_pdf_path:
        try:
            doc = SimpleDocTemplate(call_transcript_pdf_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
            elements = []
            
            elements.append(Paragraph("<b>SECURE AUDIO DECODER: CALL TRANSCRIPT</b>", title_style))
            elements.append(Spacer(1, 10))
            
            call_info = [
                [Paragraph("Filing ID:", meta_label_style), Paragraph(complaint_id, meta_val_style)],
                [Paragraph("Scam Risk Level:", meta_label_style), Paragraph(f"{risk_score}% RISK", meta_val_style)],
                [Paragraph("Transcript Source:", meta_label_style), Paragraph("Call Analysis Agent (Agent 1)", meta_val_style)]
            ]
            t = Table(call_info, colWidths=[120, 380])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#020b18')),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#ff3d00')),
                ('TOPPADDING', (0,0), (-1,-1), 5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                ('LEFTPADDING', (0,0), (-1,-1), 8),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 15))
            
            elements.append(Paragraph("<b>Speaker Transcript</b>", section_style))
            raw_transcript = report.get("transcript", "No transcript payload recorded.")
            # Highlight keywords if present
            keywords = report.get("keywords", [])
            highlighted_transcript = raw_transcript
            for kw in keywords:
                highlighted_transcript = highlighted_transcript.replace(kw, f"<b><font color='#ff3d00'>{kw.upper()}</font></b>")
                
            elements.append(Paragraph(highlighted_transcript, body_style))
            
            elements.append(Spacer(1, 10))
            elements.append(Paragraph("<b>Flagged Threat Keywords:</b>", section_style))
            elements.append(Paragraph(", ".join(keywords) if keywords else "None", body_style))
            
            doc.build(elements)
        except Exception as e:
            logger.error(f"Failed to generate Call Transcript PDF: {e}")

    # Build Evidence Summary PDF if applicable
    if not is_call and evidence_pdf_path:
        try:
            doc = SimpleDocTemplate(evidence_pdf_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
            elements = []
            
            elements.append(Paragraph("<b>CYBER SCAM EVIDENCE REPORT</b>", title_style))
            elements.append(Spacer(1, 10))
            
            # Evidence Details Table
            domain_name = report.get("domain", {}).get("name", "N/A")
            age_days = report.get("domain", {}).get("age_days", "N/A")
            ssl_valid = report.get("ssl", {}).get("valid", False)
            phishtank_status = report.get("phishtank", {}).get("known_phishing", False)
            
            evidence_data = [
                [Paragraph("Website URL:", meta_label_style), Paragraph(target_url, meta_val_style)],
                [Paragraph("Domain Age:", meta_label_style), Paragraph(f"{age_days} days old", meta_val_style)],
                [Paragraph("SSL Certificate:", meta_label_style), Paragraph("VALID" if ssl_valid else "INVALID / SELF-SIGNED", meta_val_style)],
                [Paragraph("PhishTank Status:", meta_label_style), Paragraph("FLAGGED AS PHISHING" if phishtank_status else "CLEAN / NOT FLAGGED", meta_val_style)],
                [Paragraph("Impersonation Match:", meta_label_style), Paragraph(f"{report.get('typosquat', {}).get('original_brand', 'None')} ({report.get('typosquat', {}).get('similarity', 0)}% similarity)", meta_val_style)],
            ]
            t = Table(evidence_data, colWidths=[130, 370])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#020b18')),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#00e676')),
                ('TOPPADDING', (0,0), (-1,-1), 5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                ('LEFTPADDING', (0,0), (-1,-1), 8),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 15))
            
            # Draw captured screenshot if available
            screenshot_rel = report.get("screenshot_url")
            if screenshot_rel and screenshot_rel.startswith("/static/"):
                screenshot_path = os.path.join(static_dir, screenshot_rel.replace("/static/", "", 1))
                if os.path.exists(screenshot_path):
                    elements.append(Paragraph("<b>Captured Website Evidence Screenshot</b>", section_style))
                    # Scale down the image to fit report safely (max width ~450px)
                    elements.append(Image(screenshot_path, width=400, height=250))
                    elements.append(Spacer(1, 10))
            
            elements.append(Paragraph("<b>ScamON AI Threat Explainer:</b>", section_style))
            ai_exp = report.get("ai_reasoning", {}).get("summary", "N/A")
            elements.append(Paragraph(ai_exp, body_style))
            
            doc.build(elements)
        except Exception as e:
            logger.error(f"Failed to generate Evidence Summary PDF: {e}")

    # Build DOCX Complaint Letter (Word format)
    try:
        doc_word = Document()
        doc_word.add_heading("Cybercrime Complaint Document", level=0)
        doc_word.add_paragraph(f"Complaint ID: {complaint_id}")
        doc_word.add_paragraph(f"Timestamp: {timestamp_str}")
        doc_word.add_paragraph(f"Scam Target URL/Source: {target_url}")
        
        doc_word.add_heading("Complaint Letter:", level=1)
        
        complaint_text = f"Dear Sir/Madam,\n\nI would like to report a suspected cyber scam detected by the ScamON AI Multi-Agent Cybersecurity Platform.\n\nThe target has been flagged for multiple suspicious behaviors, including typosquatting brand similarities, invalid/self-signed SSL certificates, domain registration age anomalies, and/or social engineering patterns detected within communication transcripts.\n\nRisk probability index: {risk_score}%\nThreat level classification: {threat_level}"
        
        if sms_sender and sms_message:
            complaint_text += f"\n\nSMS Phishing / SMiShing Telemetry:\n- SMS Sender: {sms_sender}\n- SMS Message Content: {sms_message}\n- SMS Risk Index: {sms_risk_score}%\n- Forensics Recommendation: {sms_recommendation}"
            
        if visual_scam:
            complaint_text += f"\n\nVisual Scam Investigation Telemetry:\n- Image Type: {visual_scam.get('image_type', 'N/A')}\n- Risk Score: {visual_scam.get('risk_score', 0)}% ({visual_scam.get('threat_level', 'LOW')})\n- OCR Text Snippet: {visual_scam.get('extracted_text', '')[:150]}...\n- Scam Category: {visual_scam.get('scam_category', 'N/A')}\n- Forensics Recommendation: {', '.join(visual_scam.get('recommendations', []))[:150]}"
            
        complaint_text += "\n\nKindly find the attached forensic investigation reports and supporting technical evidence for further evaluation and prompt action.\n\nRegards,\nScamON AI Compliance Network"
        
        doc_word.add_paragraph(complaint_text)
        doc_word.save(complaint_docx_path)
    except Exception as e:
        logger.error(f"Failed to generate docx complaint document: {e}")

    # Return generated files paths and default details
    ret = {
        "complaint_id": complaint_id,
        "subject": f"Complaint Regarding Suspected Cyber Scam: {domain_name if not is_call else 'Call Telemetry'}",
        "body": f"Dear Sir/Madam,\n\nI would like to report a suspected cyber scam detected by the ScamON AI Multi-Agent Cybersecurity Platform.\n\nThe investigation identified multiple indicators of phishing, impersonation, malicious behavior, and social engineering.\n\nKindly find the attached investigation report and supporting evidence for further action.\n\nThank you.\n\nRegards,\nScamON AI Compliance Network",
        "attachments": [
            {"name": "Complaint.pdf", "path": files["complaint_pdf"]},
            {"name": "Investigation_Report.pdf", "path": files["report_pdf"]},
            {"name": "Complaint.docx", "path": files["complaint_docx"]}
        ]
    }
    if is_call:
        ret["attachments"].append({"name": "Call_Transcript.pdf", "path": files["transcript_pdf"]})
    else:
        ret["attachments"].append({"name": "Evidence_Report.pdf", "path": files["evidence_pdf"]})

    return ret


def send_complaint_email(to_email: str, cc_email: str, subject: str, body: str, attachment_paths: List[str], static_dir: str) -> Dict[str, Any]:
    """Sends the complaint email containing PDF attachments via SMTP using Gmail App Passwords."""
    import re
    import traceback

    # 1. Fetch credentials
    gmail_user = os.getenv("EMAIL") or os.getenv("GMAIL_USER") or "your-email@gmail.com"
    gmail_app_password = os.getenv("APP_PASSWORD") or os.getenv("GMAIL_APP_PASSWORD") or "your-gmail-app-password"
    
    print(f"\n--- SMTP OUTBOUND EMAIL TRANSACTION ---")
    print(f"Sender Email (From): {gmail_user}")
    print(f"Recipient Email (To): {to_email}")
    print(f"CC Email: {cc_email if cc_email else 'None'}")
    print(f"Attachment List: {attachment_paths}")

    # 2. Validate email format
    EMAIL_REGEX = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$")
    if not EMAIL_REGEX.match(to_email):
        err_msg = f"Invalid recipient email address format: '{to_email}'"
        print(f"Validation Status: FAILED - {err_msg}")
        print(f"Delivery Status: FAILED")
        print(f"----------------------------------------\n")
        raise ValueError(err_msg)
        
    if cc_email and not EMAIL_REGEX.match(cc_email):
        err_msg = f"Invalid CC email address format: '{cc_email}'"
        print(f"Validation Status: FAILED - {err_msg}")
        print(f"Delivery Status: FAILED")
        print(f"----------------------------------------\n")
        raise ValueError(err_msg)

    print(f"Validation Status: PASSED")

    is_mock = False
    if not gmail_user or gmail_user == "your-email@gmail.com" or not gmail_app_password or gmail_app_password == "your-gmail-app-password":
        print("SMTP Credentials: NOT CONFIGURED (Simulation Mode Enabled)")
        is_mock = True
    else:
        print("SMTP Credentials: CONFIGURED")

    # Build Multi-part message
    msg = MIMEMultipart()
    msg['From'] = gmail_user
    msg['To'] = to_email
    if cc_email:
        msg['Cc'] = cc_email
    msg['Subject'] = subject
    
    msg.attach(MIMEText(body, 'plain'))
    
    # Attach documents
    for path in attachment_paths:
        abs_path = os.path.join(static_dir, path.replace("/static/", "", 1))
        if os.path.exists(abs_path):
            filename = os.path.basename(abs_path)
            try:
                with open(abs_path, 'rb') as f:
                    part = MIMEApplication(f.read(), Name=filename)
                part['Content-Disposition'] = f'attachment; filename="{filename}"'
                msg.attach(part)
            except Exception as e:
                logger.error(f"Failed to read attachment file: {filename}: {e}")
                
    # Deliver via SMTP if credentials are set
    smtp_auth_status = "NOT_AUTHENTICATED"
    if not is_mock:
        try:
            print("SMTP Connection: Connecting to smtp.gmail.com:465 via SSL...")
            server = smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=12)
            
            print(f"SMTP Authentication: Logging in as '{gmail_user}'...")
            server.login(gmail_user, gmail_app_password)
            smtp_auth_status = "SUCCESS"
            print("SMTP Authentication: SUCCESS")
            
            recipients = [to_email]
            if cc_email:
                recipients.append(cc_email)
                
            server.sendmail(gmail_user, recipients, msg.as_string())
            server.quit()
            print("Delivery Status: DELIVERED SUCCESSFULLY")
        except Exception as e:
            smtp_auth_status = "FAILED"
            print(f"SMTP Connection/Auth: FAILED")
            print(f"Delivery Status: FAILED")
            print(f"Traceback:\n{traceback.format_exc()}")
            print(f"----------------------------------------\n")
            raise RuntimeError(f"SMTP error: {str(e)}")
    else:
        smtp_auth_status = "SIMULATED"
        print("Delivery Status: SIMULATED (Offline Mock Mode)")

    print(f"----------------------------------------\n")

    # Persist the log into MongoDB Atlas database
    db = get_db()
    complaint_id = str(uuid.uuid4())
    if db is not None:
        try:
            complaint_report = {
                "complaint_id": complaint_id,
                "recipient": to_email,
                "cc": cc_email,
                "subject": subject,
                "body": body,
                "attachments": attachment_paths,
                "status": "DELIVERED" if not is_mock else "SIMULATED (Offline Mode)",
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC"),
                "created_at": datetime.now().isoformat()
            }
            db["email_scans"].insert_one(complaint_report)
            logger.info("Audit log stored successfully in email_scans collection.")
            
            # Save to unified investigations history
            from agents.history_helper import save_investigation
            save_investigation(
                agent_type="complaint",
                investigation_id=complaint_id,
                risk_score=0,
                threat_level="SAFE",
                input_data=f"FTC Report: {subject}",
                summary=f"Sent official cybercrime complaint package to {to_email}.",
                full_report=complaint_report,
                recommendation="Report successfully completed and archived."
            )
        except Exception as db_err:
            logger.warning(f"Failed to store audit log in MongoDB: {db_err}")

    return {
        "success": True,
        "message": "Complaint email sent successfully.",
        "recipient": to_email,
        "complaint_id": complaint_id,
        "delivery_mode": "SMTP" if not is_mock else "MOCK_SIMULATOR",
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")
    }
